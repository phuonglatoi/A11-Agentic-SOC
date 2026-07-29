from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


SOURCE = Path("outputs/Bao_cao_TTTN_A11_Agentic_SOC.docx")
OUTPUT = Path("outputs/Bao_cao_TTTN_A11_Agentic_SOC_Hoan_thien.docx")
FLOW_IMAGE = Path("report_assets/diagrams/so_do_2_3_luong_tong_quat_end_to_end.png")


def find_paragraph(doc: Document, exact_text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Không tìm thấy đoạn: {exact_text}")


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def insert_paragraph_after(
    anchor: Paragraph,
    text: str = "",
    style: str | None = None,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def add_labeled_paragraph(
    anchor: Paragraph,
    label: str,
    body: str,
) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, style="Normal")
    label_run = paragraph.add_run(label)
    label_run.bold = True
    paragraph.add_run(body)
    paragraph.paragraph_format.keep_together = True
    return paragraph


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def set_image_alt_text(paragraph: Paragraph, description: str) -> None:
    for element in paragraph._p.xpath(".//wp:docPr"):
        element.set("descr", description)
        element.set("title", "Sơ đồ luồng tổng quát A11 SOC")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if not FLOW_IMAGE.exists():
        raise FileNotFoundError(FLOW_IMAGE)

    doc = Document(SOURCE)

    # Nâng cấp Hình 2.3 thành sơ đồ end-to-end chi tiết và dành trọn một trang.
    heading = find_paragraph(doc, "2.6. Luồng xử lý sự kiện real-time")
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.keep_with_next = True

    image_paragraph = heading
    while image_paragraph._p.getnext() is not None:
        image_paragraph = Paragraph(image_paragraph._p.getnext(), image_paragraph._parent)
        if image_paragraph._p.xpath(".//w:drawing"):
            break
    else:
        raise ValueError("Không tìm thấy hình hiện tại của mục 2.6")

    clear_paragraph(image_paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(3)
    image_paragraph.paragraph_format.space_after = Pt(3)
    image_paragraph.paragraph_format.keep_with_next = True
    image_run = image_paragraph.add_run()
    image_run.add_picture(str(FLOW_IMAGE.resolve()), width=Inches(6.05))
    set_image_alt_text(
        image_paragraph,
        (
            "Luồng end-to-end của A11 SOC: Kali 192.168.228.128/24 tấn công qua "
            "OPNsense 192.168.228.142/24, NAT đến Ubuntu 192.168.1.100/24; "
            "telemetry được thu qua syslog, HEC, REST và Splunk, sau đó chuẩn hóa, "
            "correlation, triage, RAG/LLM, tạo incident, response và báo cáo."
        ),
    )

    caption = find_paragraph(doc, "Hình 2.3. Luồng xử lý sự kiện real-time")
    clear_paragraph(caption)
    caption.add_run(
        "Hình 2.3. Sơ đồ luồng tổng quát từ tấn công đến phản ứng và báo cáo"
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True

    source_note = find_paragraph(doc, "Nguồn: Nhóm A11 thiết kế.")
    # Có nhiều nguồn cùng nội dung; chọn nguồn ngay sau caption vừa sửa.
    source_candidate = Paragraph(caption._p.getnext(), caption._parent)
    if source_candidate.style.name == "Source Note":
        source_note = source_candidate
    clear_paragraph(source_note)
    source_note.add_run("Nguồn: Nhóm A11 thiết kế từ cấu hình lab và mã nguồn hệ thống.")
    source_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    old_explanation = (
        "Khi collector nhận payload, parser xác định định dạng và chuẩn hóa. "
        "Fingerprint được tìm trong các alert đang mở của cửa sổ 300 giây. "
        "Nếu có, event_count và last_seen được cập nhật; nếu không, alert mới được tạo. "
        "Triage luôn chạy lại khi có sự kiện để cho phép severity tăng theo ngưỡng. "
        "Enrichment và RAG bổ sung ngữ cảnh. High/Critical tạo hoặc cập nhật incident; "
        "Response Agent sinh action. Sau commit, EventBus phát thông báo SSE tới dashboard."
    )
    explanation = find_paragraph(doc, old_explanation)
    clear_paragraph(explanation)
    explanation.add_run(
        "Hình 2.3 mô tả toàn bộ tuyến dữ liệu của kịch bản thực nghiệm. Kali Linux "
        "192.168.228.128/24 thuộc subnet WAN lab 192.168.228.0/24 và sử dụng Nmap "
        "TCP SYN hoặc DIRB HTTP GET đến OPNsense WAN 192.168.228.142/24. Firewall "
        "kiểm tra trạng thái kết nối, Suricata quan sát lưu lượng và luật NAT chuyển "
        "tiếp TCP/80 đến Apache trên Ubuntu 192.168.1.100/24 thuộc LAN "
        "192.168.1.0/24. Các địa chỉ này đều thuộc dải riêng RFC1918; từ “WAN” trong "
        "đồ án biểu thị vùng WAN mô phỏng của phòng lab."
    )
    explanation.paragraph_format.keep_together = True

    second_explanation = insert_paragraph_after(explanation, style="Normal")
    second_explanation.add_run(
        "Telemetry được đưa về A11 SOC theo bốn tuyến: filterlog và EVE JSON qua "
        "remote syslog UDP/5514; Apache hoặc Windows qua REST/HEC TCP/8000; Splunk "
        "Universal Forwarder qua TLS TCP/9997 hoặc HEC HTTPS/8088; và Splunk Poller "
        "truy vấn Search REST v2 qua HTTPS/443 rồi POST vào /api/v1/ingest. Sau lớp "
        "xác thực, hệ thống parse, chuẩn hóa Common Event Schema, tạo fingerprint, "
        "loại trùng, correlation trong cửa sổ 300 giây và lưu bằng chứng gốc. "
        "Enrichment, Triage, RAG và Ollama tùy chọn tạo kết quả giải thích được. "
        "Low/Medium được giám sát trên dashboard; High/Critical tạo alert, incident, "
        "báo cáo và response proposal qua safety validation, analyst approval, "
        "executor và audit."
    )
    second_explanation.paragraph_format.keep_together = True

    # Bổ sung chiến lược dataset sau thiết kế luật triage/correlation.
    dataset_anchor = find_paragraph(doc, "Nguồn: app/agents/triage.py và knowledge/.")
    dataset_heading = insert_paragraph_after(
        dataset_anchor,
        "2.9.1. Chiến lược dataset và huấn luyện mô hình",
        "Heading 3",
    )
    dataset_heading.paragraph_format.keep_with_next = True

    dataset_intro = insert_paragraph_after(dataset_heading, style="Normal")
    dataset_intro.add_run(
        "Hệ thống hiện tại không bắt buộc phải huấn luyện một mô hình mới để vận hành. "
        "Pipeline lõi sử dụng parser, correlation, luật triage, enrichment, RAG và "
        "Local LLM tùy chọn nên vẫn phát hiện, tạo incident và lập báo cáo bằng logic "
        "xác định. Dataset nên được tích hợp trước hết cho kiểm thử ngoại tuyến, hiệu "
        "chỉnh ngưỡng, đo sai số và xây dựng nền tảng cho classifier hoặc anomaly "
        "detector trong giai đoạn mở rộng."
    )
    dataset_intro.paragraph_format.keep_together = True

    anchor = add_labeled_paragraph(
        dataset_intro,
        "Dữ liệu ưu tiên. ",
        (
            "Tập dữ liệu quan trọng nhất là log do chính lab tạo ra vì khớp trực tiếp "
            "với parser và hạ tầng của đề tài: OPNsense filterlog, Suricata EVE JSON, "
            "Apache access.log, Windows Security Event và kết quả Splunk. Mỗi kịch bản "
            "Nmap, DIRB, brute force, đăng nhập thất bại, xóa audit log và lưu lượng "
            "bình thường phải có scenario_id và thời gian bắt đầu/kết thúc rõ ràng."
        ),
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Lược đồ và nhãn. ",
        (
            "Mỗi mẫu nên giữ raw_event cùng normalized_event và các trường event_time, "
            "source_type, src_ip, dst_ip, src_port, dst_port, protocol, action, "
            "signature, asset_criticality, attack_label, severity_label, "
            "mitre_technique, analyst_verdict và scenario_id. Nhãn cuối cùng phải do "
            "analyst duyệt để tránh dùng chính dự đoán của hệ thống làm ground truth."
        ),
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Tập dữ liệu tham khảo. ",
        (
            "CICIDS2017 hoặc CSE-CIC-IDS2018 có thể bổ sung mẫu network flow; "
            "UNSW-NB15 hỗ trợ thử nghiệm nhiều họ tấn công; TON_IoT bổ sung telemetry "
            "mạng, hệ điều hành và IoT [16]–[18]. Các tập công khai không được đưa "
            "thẳng vào pipeline mà phải ánh xạ sang Common Event Schema và kiểm tra "
            "khác biệt phân phối so với mạng lab."
        ),
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Chia tập và chống rò rỉ dữ liệu. ",
        (
            "Chia train/validation/test theo scenario hoặc theo thời gian, ví dụ "
            "70%/15%/15%, thay vì tách ngẫu nhiên từng dòng log. Cách này tránh để các "
            "event gần như giống nhau của cùng một phiên tấn công xuất hiện đồng thời "
            "ở tập train và test. Dữ liệu phải được ẩn danh, version hóa, gắn checksum "
            "và lưu lineage của bước tiền xử lý."
        ),
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Đánh giá và triển khai. ",
        (
            "Mô hình phát hiện cần được đo precision, recall, F1, PR-AUC, false "
            "positive rate, latency p95, MTTD và MTTR. Trình tự an toàn là offline "
            "evaluation, shadow mode, analyst review, limited rollout rồi mới production. "
            "Dữ liệu replay không được phép kích hoạt hành động block hoặc isolate thật."
        ),
    )
    anchor = add_labeled_paragraph(
        anchor,
        "Khuyến nghị cho đề tài. ",
        (
            "Chưa nên fine-tune LLM ngay. Nên giữ RAG với playbook đã kiểm duyệt và "
            "bổ sung một classifier nhẹ trên normalized event nếu cần học máy. "
            "Fine-tune LLM chỉ phù hợp khi đã tích lũy đủ cặp incident–report được "
            "analyst phê duyệt, có bộ test độc lập và cơ chế rollback mô hình."
        ),
    )

    # Làm rõ hướng cải tiến nhất quán với mục dataset mới.
    future = find_paragraph(
        doc,
        (
            "Giai đoạn tiếp theo nên bổ sung message broker, worker pool, OpenTelemetry, "
            "metrics Prometheus, immutable audit, RBAC nhiều vai trò và secret manager. "
            "Triage cần được đánh giá trên tập dữ liệu gán nhãn, đo precision/recall và "
            "điều chỉnh ngưỡng theo từng nguồn. RAG có thể chuyển sang vector database "
            "với quy trình ký duyệt tài liệu. Response cần adapter EDR/ticket/email cụ thể, "
            "canary action và rollback. Local LLM cần kiểm thử prompt injection, structured "
            "output và model governance theo AI RMF."
        ),
    )
    future.add_run(
        " Nếu bổ sung học máy, ưu tiên classifier nhỏ trên normalized event và vận hành "
        "shadow mode trước; chỉ fine-tune LLM khi có đủ incident–report đã được analyst "
        "phê duyệt."
    )

    # Bổ sung nguồn dataset vào tài liệu tham khảo.
    references_anchor = find_paragraph(
        doc,
        "[15] Nhóm A11, Mã nguồn A11 Agentic SOC Automation, workspace của đề tài, 2026.",
    )
    references = [
        (
            "[16] Canadian Institute for Cybersecurity, CICIDS2017 Dataset. "
            "https://www.unb.ca/cic/datasets/ids-2017.html"
        ),
        (
            "[17] UNSW Canberra Cyber, UNSW-NB15 Dataset. "
            "https://research.unsw.edu.au/projects/unsw-nb15-dataset"
        ),
        (
            "[18] UNSW Canberra Cyber, TON_IoT Datasets. "
            "https://research.unsw.edu.au/projects/toniot-datasets"
        ),
    ]
    for reference in references:
        references_anchor = insert_paragraph_after(
            references_anchor,
            reference,
            "Normal",
        )

    set_update_fields_on_open(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
