from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


REPORT = Path("outputs/Bao_cao_TTTN_A11_Agentic_SOC_Local_First.docx")
FLOW_IMAGE = Path("report_assets/diagrams/so_do_2_3_luong_real_time_local_first.png")


def find_paragraph(doc: Document, exact_text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Could not find paragraph: {exact_text}")


def find_contains(doc: Document, needle: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Could not find text: {needle}")


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def insert_paragraph_after(
    anchor: Paragraph,
    text: str,
    style: str = "Normal",
) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    return paragraph


def insert_code_block_after(anchor: Paragraph, text: str) -> Paragraph:
    table = anchor._parent.add_table(rows=1, cols=1, width=Inches(6.5))
    anchor._p.addnext(table._tbl)
    cell = table.rows[0].cells[0]
    cell.text = text
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F4F7")
    cell._tc.get_or_add_tcPr().append(shading)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
            run.font.size = Pt(8.5)
    return cell.paragraphs[-1]


def set_image_alt_text(paragraph: Paragraph, description: str) -> None:
    for element in paragraph._p.xpath(".//wp:docPr"):
        element.set("descr", description)
        element.set(
            "title",
            "Sơ đồ luồng end-to-end local-first có RAG và n8n",
        )


def replace_flow_image(doc: Document) -> None:
    heading = find_paragraph(doc, "2.6. Luồng xử lý sự kiện real-time")
    image_paragraph = Paragraph(heading._p.getnext(), heading._parent)
    clear_paragraph(image_paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(3)
    image_paragraph.paragraph_format.space_after = Pt(3)
    image_paragraph.paragraph_format.keep_with_next = True
    run = image_paragraph.add_run()
    run.add_picture(str(FLOW_IMAGE.resolve()), width=Inches(6.05))
    set_image_alt_text(
        image_paragraph,
        (
            "Sơ đồ mô tả luồng từ Kali attacker qua OPNsense, Web/Windows target, "
            "telemetry trực tiếp về Ubuntu A11 SOC, xử lý SOCPipeline, RAG/Local "
            "LLM, tạo incident/report, approval gate và phản ứng qua n8n hoặc "
            "OPNsense. Splunk chỉ là nhánh quan sát bản sao log."
        ),
    )


def insert_final_explanations(doc: Document) -> None:
    if not any("n8n nhận hai webhook production" in p.text for p in doc.paragraphs):
        anchor = find_contains(doc, "Splunk được tách khỏi critical path")
        insert_paragraph_after(
            anchor,
            (
                "Nhánh n8n trong sơ đồ không thay thế lõi phát hiện mà đóng vai trò "
                "orchestration sau khi SOC đã tạo action. Workflow A11 SOC Local "
                "Automation nhận hai webhook production: /webhook/a11-soc-alert để "
                "ghi nhận cảnh báo High/Critical và /webhook/a11-soc-response để "
                "xử lý action đã được analyst phê duyệt. Mỗi execution gọi ngược "
                "/api/v1/automation/audit để lưu kết quả vào AuditEvent, nhờ đó báo "
                "cáo sự cố có đủ bằng chứng từ tấn công, phát hiện, phê duyệt đến "
                "phản ứng."
            ),
        )

    if not any("Knowledge Base hiện tại gồm tám playbook" in p.text for p in doc.paragraphs):
        anchor = find_paragraph(doc, "3.8. Triage, Enrichment và RAG Agent")
        first_body = Paragraph(anchor._p.getnext(), anchor._parent)
        inserted = insert_paragraph_after(
            first_body,
            "3.8.1. Knowledge Base và RAG agent hiện thực",
            "Heading 3",
        )
        insert_paragraph_after(
            inserted,
            (
                "Knowledge Base hiện tại gồm tám playbook Markdown đặt trong thư "
                "mục knowledge/: soc_operating_policy, web_attack_response, "
                "brute_force_response, windows_endpoint_response, "
                "suricata_network_scan_response, opnsense_containment_response, "
                "n8n_automation_workflow và local_ubuntu_operations. LocalKnowledgeBase "
                "lập chỉ mục token, tiêu đề và tag; khi SOCPipeline tạo query từ "
                "event_type, title, message và MITRE ID, RAG trả về tối đa ba đoạn "
                "playbook liên quan để Triage/LLM dùng làm ngữ cảnh. Dữ liệu không "
                "được gửi ra cloud và có thể kiểm thử qua endpoint POST "
                "/api/v1/knowledge/search."
            ),
        )

    if not any("Workflow n8n được đóng gói riêng trong repository" in p.text for p in doc.paragraphs):
        anchor = find_paragraph(doc, "3.10. Incident, report và response")
        source_note = find_contains(doc, "Nguồn: Nhóm A11 thiết kế.")
        if source_note._p.getparent() is not None:
            # Insert after the first explanatory paragraph of section 3.10.
            body = find_contains(doc, "Report Agent tạo Markdown")
        else:
            body = anchor
        heading = insert_paragraph_after(
            body,
            "3.10.1. Workflow n8n và cơ chế audit ngược",
            "Heading 3",
        )
        insert_paragraph_after(
            heading,
            (
                "Workflow n8n được đóng gói riêng trong repository tại "
                "n8n/workflows/a11_soc_local_automation.json. Khi chạy Docker Compose "
                "profile automation, n8n dùng biến A11_SOC_API_URL=http://api:8000 "
                "và A11_SOC_ADMIN_TOKEN để gọi lại API audit. Với RESPONSE_MODE=webhook, "
                "ResponseExecutor gửi payload {action_type, target, payload} tới "
                "http://n8n:5678/webhook/a11-soc-response; với notify_soc, hệ thống "
                "gửi cảnh báo tới http://n8n:5678/webhook/a11-soc-alert. Luồng này "
                "giúp demo SOAR local nhưng vẫn giữ nguyên human-in-the-loop: block_ip "
                "và isolate_host chỉ chạy sau khi analyst approve."
            ),
        )

    if not any("Trên máy Ubuntu dùng để chấm/demo" in p.text for p in doc.paragraphs):
        anchor = find_paragraph(doc, "PHỤ LỤC B. HƯỚNG DẪN CHẠY NHANH")
        heading = insert_paragraph_after(
            anchor,
            "B.1. Clone từ GitHub và chạy trên Ubuntu Server",
            "Heading 3",
        )
        body = insert_paragraph_after(
            heading,
            (
                "Trên máy Ubuntu dùng để chấm/demo, repository được clone trực tiếp "
                "từ GitHub. Sau khi đổi secret trong .env, Docker Compose dựng toàn "
                "bộ A11 SOC, PostgreSQL và n8n. Splunk không cần bật để kiểm thử "
                "luồng chính."
            ),
        )
        insert_code_block_after(
            body,
            "\n".join(
                [
                    "sudo apt update",
                    "sudo apt install -y git curl ca-certificates",
                    "curl -fsSL https://get.docker.com | sudo sh",
                    "sudo usermod -aG docker \"$USER\"",
                    "newgrp docker",
                    "git clone https://github.com/phuonglatoi/A11-Agentic-SOC.git",
                    "cd A11-Agentic-SOC",
                    "cp .env.example .env",
                    "# Sửa SOC_API_KEY, SOC_ADMIN_TOKEN, POSTGRES_PASSWORD, N8N_ENCRYPTION_KEY",
                    "docker compose --profile automation up -d --build",
                    "curl http://127.0.0.1:8000/health",
                ]
            ),
        )


def update_tables(doc: Document) -> None:
    completion = doc.tables[23]
    for row in completion.rows[1:]:
        item = row.cells[0].text.strip()
        if item == "Pipeline N8N/FastAPI":
            row.cells[1].text = (
                "FastAPI lõi; n8n workflow riêng có alert webhook, response "
                "webhook và audit callback"
            )
            row.cells[2].text = "Hoàn thành"
        elif item == "RAG Knowledge Base":
            row.cells[1].text = (
                "8 playbook Markdown; search API; truy xuất cục bộ không cần cloud"
            )
            row.cells[2].text = "Hoàn thành"

    api = doc.tables[25]
    existing = {
        row.cells[1].text.strip()
        for row in api.rows[1:]
        if len(row.cells) >= 2
    }
    additions = [
        ("GET", "/api/v1/knowledge", "Danh mục playbook RAG cục bộ"),
        ("POST", "/api/v1/knowledge/search", "Tra cứu playbook cho alert/query"),
        ("POST", "/api/v1/automation/audit", "n8n ghi kết quả automation vào audit"),
    ]
    for method, path, purpose in additions:
        if path in existing:
            continue
        row = api.add_row()
        row.cells[0].text = method
        row.cells[1].text = path
        row.cells[2].text = purpose

    quick = doc.tables[26]
    quick.rows[0].cells[0].text = "\n".join(
        [
            "# Ubuntu quick run",
            "git clone https://github.com/phuonglatoi/A11-Agentic-SOC.git",
            "cd A11-Agentic-SOC",
            "cp .env.example .env",
            "# Đổi secret trong .env; import n8n: /workflows/a11_soc_local_automation.json",
            "docker compose --profile automation up -d --build",
            "curl http://127.0.0.1:8000/health",
        ]
    )

    for table in (completion, api, quick):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.size = Pt(8.8 if table is quick else 9)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def main() -> None:
    if not REPORT.exists():
        raise FileNotFoundError(REPORT)
    if not FLOW_IMAGE.exists():
        raise FileNotFoundError(FLOW_IMAGE)

    doc = Document(REPORT)
    replace_flow_image(doc)
    insert_final_explanations(doc)
    update_tables(doc)
    set_update_fields_on_open(doc)
    doc.save(REPORT)
    print(REPORT.resolve())


if __name__ == "__main__":
    main()
