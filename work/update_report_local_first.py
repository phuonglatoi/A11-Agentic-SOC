from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


SOURCE = Path("outputs/Bao_cao_TTTN_A11_Agentic_SOC_Hoan_thien.docx")
OUTPUT = Path(
    "outputs/Bao_cao_TTTN_A11_Agentic_SOC_Local_First.docx"
)
FLOW_IMAGE = Path(
    "report_assets/diagrams/so_do_2_3_luong_real_time_local_first.png"
)


def find_paragraph(doc: Document, exact_text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Không tìm thấy đoạn: {exact_text}")


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def replace_paragraph(
    doc: Document,
    old_text: str,
    new_text: str,
) -> Paragraph:
    paragraph = find_paragraph(doc, old_text)
    clear_paragraph(paragraph)
    paragraph.add_run(new_text)
    return paragraph


def insert_paragraph_after(
    anchor: Paragraph,
    text: str,
    style: str = "Normal",
) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def set_image_alt_text(paragraph: Paragraph, description: str) -> None:
    for element in paragraph._p.xpath(".//wp:docPr"):
        element.set("descr", description)
        element.set(
            "title",
            "Luồng xử lý sự kiện real-time local-first của A11 SOC",
        )


def update_source_table(doc: Document) -> None:
    table = doc.tables[9]
    replacements = {
        "Apache": (
            "Combined access log",
            "REST/HEC trực tiếp; bản sao Splunk tùy chọn",
            "DIRB, truy cập .env, 404 burst",
        ),
        "Suricata": (
            "EVE JSON",
            "REST/HEC trực tiếp hoặc syslog",
            "Nmap scan, signature IDS",
        ),
        "Windows": (
            "Security Event JSON",
            "REST/HEC trực tiếp",
            "4625 brute force, 1102 clear log",
        ),
        "OPNsense": (
            "Syslog text/JSON",
            "Syslog UDP/5514 trực tiếp",
            "Firewall deny/pass, IDS",
        ),
    }
    for row in list(table.rows[1:]):
        key = row.cells[0].text.strip()
        if key == "Splunk":
            table._tbl.remove(row._tr)
            continue
        if key not in replacements:
            continue
        values = replacements[key]
        for index, value in enumerate(values, start=1):
            row.cells[index].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def update_environment_table(doc: Document) -> None:
    table = doc.tables[13]
    for row in table.rows[1:]:
        component = row.cells[0].text.strip()
        if component == "Máy chủ":
            row.cells[0].text = "Máy chủ mục tiêu"
            row.cells[1].text = "Ubuntu/Apache; Windows; log shipper"
            row.cells[2].text = (
                "Sinh access.log, EVE JSON, Windows Event và gửi trực tiếp về A11"
            )
        elif component == "Ứng dụng":
            row.cells[0].text = "Máy chủ A11 SOC"
            row.cells[1].text = (
                "Ubuntu Server, Docker Compose, FastAPI, SQLAlchemy"
            )
            row.cells[2].text = (
                "API/HEC/syslog, SOCPipeline, dashboard và dữ liệu vận hành"
            )
    extra = table.add_row()
    extra.cells[0].text = "Quan sát tùy chọn"
    extra.cells[1].text = "Splunk Cloud/Enterprise"
    extra.cells[2].text = (
        "Tìm kiếm và đối chiếu bản sao raw log; không cấp dữ liệu bắt buộc"
    )


def update_compose_table(doc: Document) -> None:
    table = doc.tables[15]
    for row in table.rows[1:]:
        if row.cells[0].text.strip() == "splunk-poller":
            row.cells[3].text = (
                "Tắt mặc định; chỉ dùng tích hợp mở rộng, "
                "không thuộc luồng local-first"
            )


def update_hec_auth_table(doc: Document) -> None:
    table = doc.tables[12]
    for row in table.rows[1:]:
        if row.cells[0].text.strip() == "HEC":
            row.cells[2].text = (
                "Authorization: Splunk <SOC_API_KEY> "
                "(định dạng HEC-compatible)"
            )


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if not FLOW_IMAGE.exists():
        raise FileNotFoundError(FLOW_IMAGE)

    doc = Document(SOURCE)

    architecture_heading = find_paragraph(doc, "2.4. Kiến trúc tổng thể")
    paragraph_before_architecture = Paragraph(
        architecture_heading._p.getprevious(),
        architecture_heading._parent,
    )
    if not paragraph_before_architecture.text.strip():
        clear_paragraph(paragraph_before_architecture)

    replacements = [
        (
            "– Triển khai môi trường lab VMware với Kali Linux, OPNsense, "
            "Ubuntu/Apache và Splunk Cloud.",
            "– Triển khai môi trường lab VMware với Kali Linux, OPNsense, "
            "máy chủ mục tiêu Apache/Windows và Ubuntu Server chạy A11 SOC; "
            "Splunk chỉ dùng để quan sát log tùy chọn.",
        ),
        (
            "– Xây dựng một hệ thống SOC real-time ở mức lab/pilot, có thể nhận "
            "log từ Apache, Suricata, Windows, Splunk và syslog OPNsense.",
            "– Xây dựng một hệ thống SOC real-time ở mức lab/pilot, chạy cục bộ "
            "trên Ubuntu Server và nhận trực tiếp log Apache, Suricata, Windows "
            "và syslog OPNsense; Splunk là kênh quan sát tùy chọn.",
        ),
        (
            "Đối tượng nghiên cứu là chuỗi vận hành từ telemetry đến quyết định "
            "phản ứng trong một SOC nhỏ: collector, parser, correlation, triage, "
            "enrichment, knowledge retrieval, case management, response "
            "orchestration và audit. Môi trường thực nghiệm gồm Kali Linux ở "
            "phía WAN, OPNsense làm gateway/firewall, Ubuntu chạy Apache và "
            "Agentic SOC, cùng Splunk Cloud để quan sát log. Phạm vi không bao "
            "gồm thay thế SIEM/EDR thương mại, không thực hiện cô lập endpoint "
            "thật và không tự động chặn mạng bên ngoài khi chưa được phê duyệt.",
            "Đối tượng nghiên cứu là chuỗi vận hành từ telemetry đến quyết định "
            "phản ứng trong một SOC nhỏ: collector, parser, correlation, triage, "
            "enrichment, knowledge retrieval, case management, response "
            "orchestration và audit. Môi trường thực nghiệm gồm Kali Linux ở "
            "phía WAN, OPNsense làm gateway/firewall, máy chủ mục tiêu "
            "Apache/Windows trong LAN và một Ubuntu Server riêng chạy A11 SOC. "
            "Splunk Cloud chỉ nhận bản sao telemetry để quan sát và đối chiếu. "
            "Phạm vi không bao gồm thay thế SIEM/EDR thương mại, không thực hiện "
            "cô lập endpoint thật và không tự động chặn mạng bên ngoài khi chưa "
            "được phê duyệt.",
        ),
        (
            "Bài lab ban đầu đã tạo được log từ OPNsense, Apache và Splunk Cloud "
            "nhưng chưa có lớp xử lý thống nhất sau SIEM. Analyst vẫn phải đọc "
            "log, kết luận mức độ, tra cứu và thực hiện phản ứng bằng tay. Bài "
            "toán đặt ra là xây dựng một lõi SOC cục bộ có thể nhận dữ liệu "
            "real-time từ nhiều kênh, tương quan, giải thích quyết định và tạo "
            "luồng phản ứng có kiểm soát mà không làm gián đoạn cấu trúc lab "
            "hiện có.",
            "Bài lab ban đầu đã tạo được log từ OPNsense và Apache; Splunk Cloud "
            "hỗ trợ quan sát nhưng chưa có lớp xử lý và phản ứng thống nhất. "
            "Analyst vẫn phải đọc log, kết luận mức độ, tra cứu và thực hiện "
            "phản ứng bằng tay. Bài toán đặt ra là xây dựng một lõi SOC cục bộ "
            "trên Ubuntu có thể nhận trực tiếp dữ liệu real-time từ nhiều kênh, "
            "tương quan, giải thích quyết định và tạo luồng phản ứng có kiểm "
            "soát mà không phụ thuộc Splunk.",
        ),
        (
            "Kiến trúc gồm bốn lớp. Lớp telemetry nhận dữ liệu từ "
            "OPNsense/Suricata, Apache, Windows và Splunk Cloud. Lớp ingest "
            "chuẩn hóa định dạng, tạo fingerprint và lưu bằng chứng. Lớp agent "
            "thực hiện triage, enrichment, RAG/LLM, tạo report và action "
            "proposal. Lớp vận hành cung cấp dashboard, incident queue, approval "
            "gate, response adapter và audit trail. Mỗi lớp có giao diện rõ ràng "
            "để có thể thay nguồn log hoặc mô hình AI mà không thay toàn bộ hệ "
            "thống.",
            "Kiến trúc gồm bốn lớp. Lớp telemetry nhận trực tiếp dữ liệu từ "
            "OPNsense/Suricata, Apache và Windows; Splunk nằm ngoài critical path "
            "và chỉ quan sát bản sao log. Lớp ingest trên Ubuntu chuẩn hóa định "
            "dạng, tạo fingerprint và lưu bằng chứng. Lớp agent thực hiện triage, "
            "enrichment, RAG/LLM, tạo report và action proposal. Lớp vận hành "
            "cung cấp dashboard, incident queue, approval gate, response adapter "
            "và audit trail. Mỗi lớp có giao diện rõ ràng để có thể thay nguồn "
            "log hoặc mô hình AI mà không thay toàn bộ hệ thống.",
        ),
        (
            "Kali Linux đóng vai trò nguồn kiểm thử bên WAN. OPNsense có các "
            "interface WAN, LAN và DMZ, thực hiện NAT và firewall. Ubuntu chạy "
            "Apache làm mục tiêu, đồng thời có thể chạy A11 SOC. Splunk Universal "
            "Forwarder gửi access log lên Splunk Cloud; song song, shippers hoặc "
            "poller đưa dữ liệu vào collector cục bộ. Syslog từ OPNsense được "
            "giới hạn tới UDP 5514 trong mạng lab.",
            "Kali Linux đóng vai trò nguồn kiểm thử bên WAN. OPNsense có các "
            "interface WAN, LAN và DMZ, thực hiện NAT và firewall. Máy chủ "
            "192.168.1.100 chạy Apache/Suricata làm mục tiêu; máy Windows "
            "192.168.1.101 cung cấp Security Event; Ubuntu Server "
            "192.168.1.10 chạy A11 SOC độc lập. Log shipper gửi access.log và "
            "EVE JSON trực tiếp qua REST/HEC TCP/8000, Windows collector dùng "
            "cùng kênh, còn OPNsense gửi syslog UDP/5514. Splunk chỉ nhận bản "
            "sao telemetry để quan sát và không nằm trên tuyến xử lý bắt buộc.",
        ),
        (
            "Hình 2.3 mô tả toàn bộ tuyến dữ liệu của kịch bản thực nghiệm. Kali "
            "Linux 192.168.228.128/24 thuộc subnet WAN lab 192.168.228.0/24 và "
            "sử dụng Nmap TCP SYN hoặc DIRB HTTP GET đến OPNsense WAN "
            "192.168.228.142/24. Firewall kiểm tra trạng thái kết nối, Suricata "
            "quan sát lưu lượng và luật NAT chuyển tiếp TCP/80 đến Apache trên "
            "Ubuntu 192.168.1.100/24 thuộc LAN 192.168.1.0/24. Các địa chỉ này "
            "đều thuộc dải riêng RFC1918; từ “WAN” trong đồ án biểu thị vùng WAN "
            "mô phỏng của phòng lab.",
            "Hình 2.3 mô tả toàn bộ tuyến dữ liệu local-first của kịch bản thực "
            "nghiệm. Kali Linux 192.168.228.128/24 thuộc subnet WAN lab "
            "192.168.228.0/24 và sử dụng Nmap TCP SYN hoặc DIRB HTTP GET đến "
            "OPNsense WAN 192.168.228.142/24. Firewall kiểm tra trạng thái kết "
            "nối, Suricata quan sát lưu lượng và luật NAT chuyển tiếp TCP/80 đến "
            "Apache 192.168.1.100/24 trong LAN 192.168.1.0/24. Ubuntu A11 SOC "
            "192.168.1.10/24 là máy chủ độc lập, không đồng thời đóng vai trò "
            "web target. Các địa chỉ này đều thuộc dải riêng RFC1918; từ “WAN” "
            "biểu thị vùng WAN mô phỏng của phòng lab.",
        ),
        (
            "Telemetry được đưa về A11 SOC theo bốn tuyến: filterlog và EVE JSON "
            "qua remote syslog UDP/5514; Apache hoặc Windows qua REST/HEC "
            "TCP/8000; Splunk Universal Forwarder qua TLS TCP/9997 hoặc HEC "
            "HTTPS/8088; và Splunk Poller truy vấn Search REST v2 qua HTTPS/443 "
            "rồi POST vào /api/v1/ingest. Sau lớp xác thực, hệ thống parse, chuẩn "
            "hóa Common Event Schema, tạo fingerprint, loại trùng, correlation "
            "trong cửa sổ 300 giây và lưu bằng chứng gốc. Enrichment, Triage, "
            "RAG và Ollama tùy chọn tạo kết quả giải thích được. Low/Medium được "
            "giám sát trên dashboard; High/Critical tạo alert, incident, báo cáo "
            "và response proposal qua safety validation, analyst approval, "
            "executor và audit.",
            "Trong tuyến chính, OPNsense gửi filterlog qua syslog UDP/5514; log "
            "shipper trên Apache/Suricata gửi access.log và EVE JSON qua "
            "REST/HEC TCP/8000; Windows collector gửi Security Event JSON qua "
            "REST/HEC TCP/8000. Container api trên Ubuntu tiếp nhận dữ liệu và "
            "gọi SOCPipeline.process(). Bước 3 lần lượt xác thực, parse, chuẩn "
            "hóa Common Event Schema, tạo fingerprint, tương quan trong cửa sổ "
            "300 giây và lưu raw evidence, SecurityEvent, Alert cùng AuditEvent. "
            "Dedup không xóa log gốc mà chỉ ngăn tạo quá nhiều Alert.",
        ),
        (
            "Dữ liệu ưu tiên. Tập dữ liệu quan trọng nhất là log do chính lab tạo "
            "ra vì khớp trực tiếp với parser và hạ tầng của đề tài: OPNsense "
            "filterlog, Suricata EVE JSON, Apache access.log, Windows Security "
            "Event và kết quả Splunk. Mỗi kịch bản Nmap, DIRB, brute force, đăng "
            "nhập thất bại, xóa audit log và lưu lượng bình thường phải có "
            "scenario_id và thời gian bắt đầu/kết thúc rõ ràng.",
            "Dữ liệu ưu tiên. Tập dữ liệu quan trọng nhất là log do chính lab tạo "
            "ra vì khớp trực tiếp với parser và hạ tầng của đề tài: OPNsense "
            "filterlog, Suricata EVE JSON, Apache access.log và Windows Security "
            "Event. Kết quả hiển thị trên Splunk chỉ dùng đối chiếu bản sao "
            "telemetry, không được xem là nguồn ground truth riêng. Mỗi kịch bản "
            "Nmap, DIRB, brute force, đăng nhập thất bại, xóa audit log và lưu "
            "lượng bình thường phải có scenario_id và thời gian bắt đầu/kết thúc "
            "rõ ràng.",
        ),
        (
            "Apache được cài trên Ubuntu và access.log là nguồn bằng chứng chính "
            "cho kịch bản web. Port forward tại OPNsense chuyển WAN:80 tới địa "
            "chỉ Ubuntu trong LAN. Rule liên quan được kiểm tra để tránh mở dịch "
            "vụ ngoài phạm vi lab. Universal Forwarder theo dõi "
            "/var/log/apache2/access.log và gửi dữ liệu lên Splunk Cloud.",
            "Apache được cài trên máy chủ mục tiêu Ubuntu 192.168.1.100 và "
            "access.log là nguồn bằng chứng chính cho kịch bản web. Port forward "
            "tại OPNsense chuyển WAN:80 tới địa chỉ này trong LAN. Log shipper "
            "theo dõi /var/log/apache2/access.log và gửi sự kiện trực tiếp tới "
            "A11 SOC 192.168.1.10 qua REST/HEC TCP/8000. Universal Forwarder có "
            "thể gửi một bản sao lên Splunk Cloud để quan sát, nhưng không thuộc "
            "luồng phát hiện bắt buộc.",
        ),
        (
            "Kali chạy DIRB tới địa chỉ WAN của OPNsense; port forward chuyển "
            "request tới Apache. Số lượng access log tăng nhanh và chứa "
            "user-agent/đường dẫn đặc trưng. Splunk Cloud nhận dữ liệu qua "
            "Universal Forwarder và cho phép quan sát xu hướng. A11 SOC có thể "
            "nhận cùng log trực tiếp qua HEC để giảm độ trễ, hoặc poll kết quả từ "
            "endpoint v2 /services/search/v2/jobs/export khi tài khoản Splunk "
            "Cloud được cấp REST API.",
            "Kali chạy DIRB tới địa chỉ WAN của OPNsense; port forward chuyển "
            "request tới Apache. Số lượng access log tăng nhanh và chứa "
            "user-agent/đường dẫn đặc trưng. Log shipper gửi từng sự kiện trực "
            "tiếp vào A11 SOC qua REST/HEC để bảo đảm độ trễ thấp. Splunk Cloud "
            "nhận bản sao qua Universal Forwarder và chỉ dùng quan sát xu hướng, "
            "tìm kiếm raw log và đối chiếu với Alert trên dashboard A11. Profile "
            "splunk-poller tắt mặc định; chỉ bật khi cần thử nghiệm tích hợp REST "
            "với một SIEM có sẵn.",
        ),
        (
            "Chương 3 đã mô tả đầy đủ từ hạ tầng mạng đến pipeline ứng dụng, "
            "agent, dashboard và kiểm soát triển khai. Hệ thống giữ lại giá trị "
            "của OPNsense/Splunk hiện có đồng thời bổ sung lớp automation cục bộ. "
            "Chương 4 xác minh các tuyên bố bằng test và số liệu được sinh trực "
            "tiếp từ mã nguồn.",
            "Chương 3 đã mô tả đầy đủ từ hạ tầng mạng đến pipeline ứng dụng, "
            "agent, dashboard và kiểm soát triển khai. A11 SOC chạy độc lập trên "
            "Ubuntu, nhận trực tiếp telemetry từ các nguồn local; OPNsense cung "
            "cấp firewall/syslog và Splunk chỉ là công cụ quan sát tùy chọn. "
            "Chương 4 xác minh các tuyên bố bằng test và số liệu được sinh trực "
            "tiếp từ mã nguồn.",
        ),
        (
            "Số liệu trên phản ánh chi phí pipeline trong tiến trình, không phải "
            "khả năng tải của một cụm production. Latency thực tế còn phụ thuộc "
            "network, PostgreSQL, số client, kích thước log, truy vấn Splunk và "
            "model Ollama. Tuy vậy, kết quả đủ chứng minh thiết kế đáp ứng quy "
            "mô lab theo thời gian thực và correlation không làm mất bằng chứng.",
            "Số liệu trên phản ánh chi phí pipeline trong tiến trình, không phải "
            "khả năng tải của một cụm production. Latency tuyến chính phụ thuộc "
            "network nội bộ, PostgreSQL, số client, kích thước log và model "
            "Ollama; truy vấn Splunk không ảnh hưởng vì nằm ngoài critical path. "
            "Kết quả đủ chứng minh thiết kế đáp ứng quy mô lab theo thời gian "
            "thực và correlation không làm mất bằng chứng.",
        ),
        (
            "– Splunk Cloud REST phụ thuộc gói dịch vụ/quyền; luồng song song qua "
            "HEC là phương án thực tế hơn trong lab.",
            "– Splunk Cloud REST phụ thuộc gói dịch vụ/quyền nên chỉ được giữ "
            "như nhánh quan sát hoặc tích hợp mở rộng; luồng local trực tiếp vào "
            "A11 qua REST/HEC/syslog là tuyến vận hành chính.",
        ),
        (
            "Đề tài đã hoàn thành một hệ thống Agentic SOC real-time ở mức "
            "lab/pilot, kết hợp OPNsense, Apache, Splunk với lớp automation cục "
            "bộ. Sản phẩm nhận log qua REST, HEC và syslog; chuẩn hóa bốn nhóm "
            "định dạng; lưu raw evidence; tương quan; triage; enrichment; RAG; "
            "Local LLM tùy chọn; incident; report; response proposal; approval; "
            "audit và dashboard SSE. Docker Compose cung cấp phương án triển "
            "khai tối thiểu và các profile AI, workflow, Splunk.",
            "Đề tài đã hoàn thành một hệ thống Agentic SOC real-time ở mức "
            "lab/pilot, chạy local-first trên Ubuntu Server và nhận trực tiếp "
            "telemetry từ OPNsense, Apache/Suricata và Windows. Sản phẩm nhận log "
            "qua REST, HEC-compatible API và syslog; chuẩn hóa bốn nhóm định "
            "dạng; lưu raw evidence; tương quan; triage; enrichment; RAG; Local "
            "LLM tùy chọn; incident; report; response proposal; approval; audit "
            "và dashboard SSE. Docker Compose cung cấp phương án triển khai tối "
            "thiểu; Splunk được giữ trong profile tùy chọn để quan sát và tích "
            "hợp mở rộng, không phải điều kiện vận hành của A11 SOC.",
        ),
    ]
    for old_text, new_text in replacements:
        replace_paragraph(doc, old_text, new_text)

    heading = find_paragraph(doc, "2.6. Luồng xử lý sự kiện real-time")
    image_paragraph = Paragraph(heading._p.getnext(), heading._parent)
    if not image_paragraph._p.xpath(".//w:drawing"):
        raise ValueError("Không tìm thấy hình ngay sau mục 2.6")
    clear_paragraph(image_paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(3)
    image_paragraph.paragraph_format.space_after = Pt(3)
    image_paragraph.paragraph_format.keep_with_next = True
    image_run = image_paragraph.add_run()
    image_run.add_picture(str(FLOW_IMAGE.resolve()), width=Inches(6.05))
    set_image_alt_text(
        image_paragraph,
        (
            "Sơ đồ local-first: Kali tấn công qua OPNsense đến Apache; "
            "OPNsense, Apache/Suricata và Windows gửi telemetry trực tiếp về "
            "Ubuntu A11 SOC. SOCPipeline thực hiện auth, parse, normalize, "
            "dedup, correlate và persist rồi bàn giao Correlated Alert Context "
            "cho Agentic AI. Splunk nhận bản sao log để quan sát tùy chọn."
        ),
    )

    pipeline_paragraph = find_paragraph(
        doc,
        (
            "Trong tuyến chính, OPNsense gửi filterlog qua syslog UDP/5514; log "
            "shipper trên Apache/Suricata gửi access.log và EVE JSON qua "
            "REST/HEC TCP/8000; Windows collector gửi Security Event JSON qua "
            "REST/HEC TCP/8000. Container api trên Ubuntu tiếp nhận dữ liệu và "
            "gọi SOCPipeline.process(). Bước 3 lần lượt xác thực, parse, chuẩn "
            "hóa Common Event Schema, tạo fingerprint, tương quan trong cửa sổ "
            "300 giây và lưu raw evidence, SecurityEvent, Alert cùng AuditEvent. "
            "Dedup không xóa log gốc mà chỉ ngăn tạo quá nhiều Alert."
        ),
    )
    handoff = insert_paragraph_after(
        pipeline_paragraph,
        (
            "SOCPipeline là thành phần điều phối đại diện cho bước 3. Đầu ra của "
            "bước này là Correlated Alert Context gồm alert_id, fingerprint, "
            "event_count, first_seen, last_seen, normalized_event và tham chiếu "
            "bằng chứng gốc. Trong hiện thực hiện tại, SOCPipeline bàn giao context "
            "cho bước 4 bằng lời gọi hàm nội bộ tới Enrichment Agent, Triage, RAG "
            "và Local LLM tùy chọn; không gọi API cloud. Sau khi tạo Incident, "
            "ResponseAction và báo cáo, transaction được commit rồi EventBus phát "
            "Alert/Incident qua SSE tới dashboard.",
        ),
    )
    handoff.paragraph_format.keep_together = True
    splunk_note = insert_paragraph_after(
        handoff,
        (
            "Splunk được tách khỏi critical path. Universal Forwarder hoặc "
            "collector có thể gửi một bản sao raw telemetry lên Splunk "
            "Cloud/Enterprise để tìm kiếm, lập biểu đồ và đối chiếu với A11. "
            "A11 vẫn thu nhận, phát hiện, tạo Incident và báo cáo đầy đủ khi "
            "Splunk tắt hoặc mất kết nối; profile splunk-poller chỉ là adapter "
            "mở rộng và mặc định không được khởi động."
        ),
    )
    splunk_note.paragraph_format.keep_together = True

    update_source_table(doc)
    source_note = find_paragraph(doc, "Nguồn: Mã nguồn parser và tài liệu lab.")
    source_note.paragraph_format.keep_with_next = False
    source_note.paragraph_format.keep_together = True
    update_environment_table(doc)
    update_compose_table(doc)
    update_hec_auth_table(doc)

    set_update_fields_on_open(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
