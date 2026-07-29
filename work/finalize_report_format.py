from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


REPORT = Path("outputs/Bao_cao_TTTN_A11_Agentic_SOC_Hoan_thien.docx")


def main() -> None:
    doc = Document(REPORT)

    # Giữ mục lục trong hai trang, tránh vừa quá dày vừa phát sinh trang thứ ba chỉ có hai dòng.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.lower().startswith("toc "):
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(18)
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # Tránh giãn trắng bất thường ở các tài liệu tham khảo có URL dài.
    in_references = False
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "TÀI LIỆU THAM KHẢO":
            in_references = True
            continue
        if in_references and paragraph.text.strip().startswith("["):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Mục lục đã được Word cập nhật; không yêu cầu cập nhật lại khi mở.
    update_fields = doc.settings._element.find(qn("w:updateFields"))
    if update_fields is not None:
        update_fields.set(qn("w:val"), "false")

    doc.save(REPORT)
    print(REPORT.resolve())


if __name__ == "__main__":
    main()
