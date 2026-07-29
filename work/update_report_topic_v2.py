from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


REPORT = Path("outputs/Bao_cao_TTTN_A11_Agentic_SOC_Local_First.docx")

NEW_TITLE_UPPER = (
    "XÂY DỰNG HỆ THỐNG SOC REAL-TIME LOCAL-FIRST\n"
    "ỨNG DỤNG AGENTIC AI, RAG VÀ N8N TRÊN UBUNTU SERVER"
)
NEW_TITLE_SENTENCE = (
    "Xây dựng hệ thống SOC real-time local-first ứng dụng Agentic AI, "
    "RAG và n8n trên Ubuntu Server"
)

DIAGRAMS = {
    "Hình 2.1.": (
        "Hình 2.1. Kiến trúc tổng thể A11 Agentic SOC local-first",
        Path("report_assets/diagrams/so_do_2_1_kien_truc_tong_the.png"),
        "Kiến trúc tổng thể A11 SOC local-first: telemetry, collector, SOCPipeline, Agentic AI/RAG, operations, data và Splunk quan sát.",
    ),
    "Hình 2.2.": (
        "Hình 2.2. Mô hình mạng lab, IP và tuyến telemetry local-first",
        Path("report_assets/diagrams/so_do_2_2_mo_hinh_mang_lab.png"),
        "Mô hình mạng lab nêu rõ IP Kali, OPNsense WAN/LAN, máy Web/Windows, A11 SOC Ubuntu và các giao thức REST/HEC/syslog.",
    ),
    "Hình 2.3.": (
        "Hình 2.3. Sơ đồ luồng tổng quát từ tấn công đến phản ứng và báo cáo",
        Path("report_assets/diagrams/so_do_2_3_luong_real_time_local_first.png"),
        "Luồng end-to-end từ tấn công, sinh log, ingest local, xử lý real-time, RAG, incident, approval, n8n và báo cáo.",
    ),
    "Hình 2.4.": (
        "Hình 2.4. Mô hình phối hợp SOCPipeline, RAG và n8n",
        Path("report_assets/diagrams/so_do_2_4_phoi_hop_agent.png"),
        "Mô hình phối hợp các agent trong SOCPipeline, RAG knowledge base, local LLM, approval gate, n8n và audit.",
    ),
    "Hình 3.10.": (
        "Hình 3.10. Docker Compose trên Ubuntu Server với n8n workflow",
        Path("report_assets/diagrams/so_do_3_1_trien_khai_docker.png"),
        "Kiến trúc triển khai Docker Compose trên Ubuntu Server gồm api, postgres, n8n, ollama tùy chọn, Splunk quan sát tùy chọn.",
    ),
    "Hình 3.11.": (
        "Hình 3.11. Sequence từ tấn công đến RAG, approval, n8n và audit",
        Path("report_assets/diagrams/so_do_3_2_sequence_phan_ung.png"),
        "Sequence xử lý sự kiện từ attacker đến collector, SOCPipeline, RAG/agent, analyst approval, n8n execution và audit callback.",
    ),
    "Hình 3.12.": (
        "Hình 3.12. Vòng đời alert, incident, response action và audit",
        Path("report_assets/diagrams/so_do_3_3_vong_doi_alert.png"),
        "Vòng đời dữ liệu từ raw event đến alert, incident, response action, execution, audit event, report và dashboard SSE.",
    ),
}


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, *, bold: bool = False, size: float | None = None) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if bold:
        run.bold = True
    if size is not None:
        run.font.size = Pt(size)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
        paragraph._p = paragraph._element = None


def find_caption(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Figure Caption" and paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Không tìm thấy caption {prefix}")


def previous_image_paragraph(caption: Paragraph) -> Paragraph:
    node = caption._p.getprevious()
    while node is not None:
        paragraph = Paragraph(node, caption._parent)
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
        if paragraph.text.strip():
            break
        node = node.getprevious()
    raise ValueError(f"Không tìm thấy paragraph ảnh trước caption: {caption.text}")


def set_image_alt_text(paragraph: Paragraph, description: str) -> None:
    for element in paragraph._p.xpath(".//wp:docPr"):
        element.set("descr", description)
        element.set("title", description[:240])


def replace_diagrams(doc: Document) -> None:
    for prefix, (caption_text, image_path, alt_text) in DIAGRAMS.items():
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        caption = find_caption(doc, prefix)
        image_paragraph = previous_image_paragraph(caption)
        clear_paragraph(image_paragraph)
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.space_before = Pt(3)
        image_paragraph.paragraph_format.space_after = Pt(3)
        image_paragraph.paragraph_format.keep_with_next = True
        image_paragraph.add_run().add_picture(str(image_path.resolve()), width=Inches(6.05))
        set_image_alt_text(image_paragraph, alt_text)
        set_paragraph_text(caption, caption_text, bold=False, size=10)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def update_title(doc: Document) -> None:
    old_title = "XÂY DỰNG HỆ THỐNG TỰ ĐỘNG HÓA VẬN HÀNH SOC BẰNG MÔ HÌNH AGENTIC AI TRÊN HẠ TẦNG CỤC BỘ"
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in {old_title, NEW_TITLE_UPPER.replace("\n", " ")} or "TỰ ĐỘNG HÓA VẬN HÀNH SOC" in paragraph.text:
            clear_paragraph(paragraph)
            first, second = NEW_TITLE_UPPER.split("\n", 1)
            run = paragraph.add_run(first)
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run.font.size = Pt(14)
            run.add_break()
            run2 = paragraph.add_run(second)
            run2.bold = True
            run2.font.name = "Times New Roman"
            run2._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            run2._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run2.font.size = Pt(14)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == "Tên đề tài":
                row.cells[1].text = NEW_TITLE_SENTENCE
                for p in row.cells[1].paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                        run.font.size = Pt(11)


def update_context_paragraphs(doc: Document) -> None:
    replacements = {
        "Kiến trúc gồm bốn lớp.": (
            "Hình 2.1 trình bày kiến trúc local-first của A11 SOC theo đề tài mới. "
            "Telemetry từ OPNsense/Suricata, Apache và Windows được đưa trực tiếp vào "
            "A11 Collector trên Ubuntu Server; SOCPipeline thực hiện parse, normalize, "
            "dedup, correlate và persist; Agentic AI/RAG sinh nhận định, khuyến nghị, "
            "incident và báo cáo. Splunk/SIEM không nằm trên tuyến xử lý chính mà chỉ "
            "nhận bản sao log để quan sát, tìm kiếm raw log và đối chiếu khi demo."
        ),
        "Kali Linux đóng vai trò nguồn kiểm thử bên WAN.": (
            "Hình 2.2 mô tả rõ mạng lab và địa chỉ IP: Kali 192.168.228.128/24 ở WAN lab "
            "tấn công qua OPNsense WAN 192.168.228.142/24; OPNsense chuyển tiếp vào LAN "
            "192.168.1.0/24, nơi có Web/Suricata 192.168.1.100/24, Windows endpoint "
            "192.168.1.101/24 và A11 SOC Ubuntu 192.168.1.10/24. Telemetry chính đi bằng "
            "REST/HEC TCP/8000 và syslog UDP/5514; Splunk chỉ là nhánh nhận bản sao."
        ),
        "Orchestrator được hiện thực trong SOCPipeline.": (
            "Hình 2.4 cho thấy SOCPipeline là điều phối trung tâm: sau khi sự kiện được "
            "chuẩn hóa, pipeline gọi Triage, Enrichment, RAG Agent và Local LLM/fallback "
            "theo thứ tự xác định. Đầu ra là alert/incident/report và response action dạng "
            "proposal. Với hành động rủi ro cao như block_ip hoặc isolate_host, hệ thống "
            "bắt buộc approval trước khi n8n hoặc adapter OPNsense thực thi."
        ),
        "Compose mặc định khởi động api và postgres.": (
            "Hình 3.10 thể hiện triển khai chuẩn khi clone repository về Ubuntu Server: "
            "Docker Compose khởi động api và postgres theo mặc định; n8n nằm trong profile "
            "automation; Ollama nằm trong profile ai; Splunk/poller là nhánh quan sát tùy chọn. "
            "API publish HTTP theo SOC_HTTP_BIND, syslog UDP/5514 nhận log cục bộ, data và "
            "knowledge được mount vào container để bảo toàn bằng chứng và playbook."
        ),
        "Report Agent tạo Markdown gồm tóm tắt": (
            "Hình 3.11 mô tả sequence phản ứng: attacker tạo Nmap/DIRB/brute force; Web, "
            "Windows hoặc OPNsense sinh log; Collector xác thực và chuyển sự kiện vào "
            "SOCPipeline; pipeline dedup/correlate trong cửa sổ 300 giây, truy xuất RAG "
            "playbook, sinh severity/MITRE/recommendation, rồi tạo incident và action. "
            "Analyst phê duyệt trước khi n8n/adapter thực thi; kết quả được ghi ngược vào "
            "AuditEvent để Report Agent tổng hợp timeline, bằng chứng và khuyến nghị."
        ),
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for prefix, replacement in replacements.items():
            if text.startswith(prefix):
                set_paragraph_text(paragraph, replacement, size=11)
                paragraph.paragraph_format.first_line_indent = Inches(0.3)
                paragraph.paragraph_format.line_spacing = 1.15

    rag_old = "RAG kết hợp truy xuất tài liệu với sinh nội dung. Knowledge Base của đề tài nạp bốn tài liệu Markdown"
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(rag_old):
            set_paragraph_text(
                paragraph,
                (
                    "RAG kết hợp truy xuất tài liệu với sinh nội dung. Knowledge Base của đề tài "
                    "hiện nạp tám playbook Markdown trong thư mục knowledge/: chính sách vận hành "
                    "SOC, phản ứng web attack, brute force, Windows endpoint, Suricata network scan, "
                    "OPNsense containment, n8n automation workflow và vận hành Ubuntu local. "
                    "Truy xuất dùng chỉ mục cục bộ theo token/tag, không gửi dữ liệu ra cloud; "
                    "kết quả RAG cung cấp ngữ cảnh cho Triage/LLM và Report Agent."
                ),
                size=11,
            )
            paragraph.paragraph_format.first_line_indent = Inches(0.3)
            paragraph.paragraph_format.line_spacing = 1.15


def remove_duplicate_n8n_paragraph(doc: Document) -> None:
    seen = False
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith("Nhánh n8n trong sơ đồ không thay thế lõi phát hiện"):
            if seen:
                delete_paragraph(paragraph)
            else:
                seen = True


def update_splunk_caption(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Figure Caption" and paragraph.text.strip().startswith("Hình 3.15."):
            set_paragraph_text(
                paragraph,
                "Hình 3.15. Splunk/SIEM quan sát bản sao Apache access log",
                size=10,
            )
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def main() -> None:
    if not REPORT.exists():
        raise FileNotFoundError(REPORT)

    doc = Document(REPORT)
    update_title(doc)
    replace_diagrams(doc)
    update_context_paragraphs(doc)
    remove_duplicate_n8n_paragraph(doc)
    update_splunk_caption(doc)
    set_update_fields_on_open(doc)
    doc.save(REPORT)
    print(REPORT.resolve())


if __name__ == "__main__":
    main()
